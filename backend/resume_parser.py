import pdfplumber
import docx
import io


def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:
    """Parse resume from bytes (PDF or DOCX) and return raw text."""
    text = ""

    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    elif filename.lower().endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    elif filename.lower().endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")

    else:
        # Try to decode as plain text
        text = file_bytes.decode("utf-8", errors="ignore")

    return text.strip()


def parse_resume_file(file_path: str) -> str:
    """Parse resume from a file path and return raw text."""
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    filename = file_path.split("/")[-1]
    return parse_resume_bytes(file_bytes, filename)
